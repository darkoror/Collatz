from operator import itemgetter

from docx import Document
from docx.shared import RGBColor


def result(lowest_number, highest_number: int, quantity_to_return: int, filename='demo'):
    all_results = {}  # {number: collatz_list}
    all_numb_and_len = {}  # {number: len(collatz_list)
    all_len = []  # list of lens of each collatz_list  # we need it in generator

    for num in range(lowest_number, highest_number + 1):
        collatz_list = collatz(num)

        # remove even numbers
        collatz_list = [i for i in collatz_list if i % 2 == 1]

        all_results[num] = collatz_list
        all_numb_and_len[num] = len(collatz_list)
        all_len.append(len(collatz_list))

    # sort by len of collatz_list (from highest to lowest len)
    sorted_numb_and_len = sorted(all_numb_and_len.items(), key=itemgetter(1), reverse=True)

    exclude_repeated_even_numbers_generator = exclude_repeated_even_numbers(sorted_numb_and_len, all_len)
    top_list = []
    for i in exclude_repeated_even_numbers_generator:
        top_list.append(i)
        if len(top_list) == quantity_to_return:
            break

    # print result into console
    # print_results(top_list, all_results)

    # write results to MS Word file
    write_results_to_word(filename, top_list, all_results)


def write_results_to_word(filename, top_list, all_results):
    document = Document()

    for index, value in list(enumerate(top_list)):
        number, length = value
        p = document.add_paragraph('')
        p.add_run(f"{index + 1}. {number} ({length}): [").bold = True
        for i in all_results[number]:
            if i % 3 == 0:
                p.add_run(f'{i}, ').font.color.rgb = RGBColor(255, 0, 0)
            else:
                p.add_run(f'{i}, ')
        p.add_run(']')

    document.save(f'{filename}.docx')


def print_results(top_list, all_results):
    for index, value in list(enumerate(top_list)):
        number, length = value
        print(
            f"{index + 1}. {number} ({length}): {all_results[number]}\n"
        )


def exclude_repeated_even_numbers(sorted_list, all_len):
    # generator, which skips values where collatz_list_len (in this case - sorted_list[1])
    # and number is even (sorted_list[0])
    for numb, numb_len in sorted_list:
        if all_len.count(numb_len) > 1 and numb % 2 == 0:
            continue
        yield numb, numb_len


def collatz(n):
    sp = [n]

    if n < 1:
        return []

    while n > 1:
        if n % 2 == 0:
            n = n // 2
        else:
            n = 3 * n + 1
        sp.append(n)

    return sp


if __name__ == '__main__':
    result(3, 100, 10)
